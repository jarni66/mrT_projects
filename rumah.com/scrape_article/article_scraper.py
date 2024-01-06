from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
import time
from bs4 import BeautifulSoup
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
import json
from docx import Document
from htmldocx import HtmlToDocx
import requests
from docx.shared import Inches
import undetected_chromedriver as uc 
import re
import os


def to_doc(x):

    if x.name == 'div':
        try:
            if 'p' in x.attrs['class']:
                document.add_paragraph(x.text)
        except:
            pass
        
    elif x.name == 'ol':
        li = x.find_all('li')
        for j in li:
            document.add_paragraph(j.text,style='List Bullet')
    elif x.name == 'h2':
        document.add_heading(x.text, level=2)
        
    elif x.name == 'figure':
        try:
            link_img = x.find('img').get('src')
            with open('temp_img/file.png', 'wb') as f:
                f.write(requests.get(link_img).content)
            document.add_picture('temp_img/file.png',width=Inches(5))

        except:
            pass




chrome_options = uc.ChromeOptions() 
chrome_options.headless = False  # Set headless to False to run in non-headless mode

driver = uc.Chrome(options=chrome_options)
with open('output_docs/done_link.json', encoding='utf-8') as f:
    done1 = json.load(f)
with open('output_docs/done_link_2.json', encoding='utf-8') as f:
    done2 = json.load(f)

done_link_1 = [i['link'] for i in done1]
done_link_2 = [i['link'] for i in done2]

done_link = done_link_1 + done_link_2
## JSON data article
# import requests
# import json

# url = "https://api-angel-green.propertyguru.com/graphql/pg"

# payload = "{\"query\":\"query Query($page: Int, $pageSize: Int, $market: String!, $language: String, $category: [String], $tag: String, $exclude: [Int], $sponsorSlug: [String]) {\\n  articleList(\\n    page: $page\\n    pageSize: $pageSize\\n    market: $market\\n    language: $language\\n    category: $category\\n    tag: $tag\\n    exclude: $exclude\\n    sponsorSlug: $sponsorSlug\\n  ) {\\n    totalCount\\n    totalPage\\n    items {\\n      id\\n      title\\n      excerpt\\n      slug\\n      link\\n      featuredImage\\n      postDate\\n      modifiedDate\\n      location\\n      author {\\n        id\\n        name\\n        slug\\n        link\\n        profilePhoto\\n        __typename\\n      }\\n      sponsor {\\n        slug\\n        name\\n        picture\\n        bio\\n        __typename\\n      }\\n      category {\\n        id\\n        name\\n        slug\\n        link\\n        __typename\\n      }\\n      tags {\\n        id\\n        name\\n        slug\\n        link\\n        __typename\\n      }\\n      profiles {\\n        id\\n        name\\n        slug\\n        link\\n        __typename\\n      }\\n      __typename\\n    }\\n    __typename\\n  }\\n}\",\"variables\":{\"page\":1,\"pageSize\":3445,\"market\":\"id\",\"language\":\"id\",\"exclude\":[84459,89258,89234,89228]}}"
# headers = {
#   'authority': 'api-angel-green.propertyguru.com',
#   'accept': '*/*',
#   'accept-language': 'en-US,en;q=0.9',
#   'cache-control': 'no-cache',
#   'content-type': 'application/json',
#   'origin': 'https://www.rumah.com',
#   'pragma': 'no-cache',
#   'referer': 'https://www.rumah.com/',
#   'sec-ch-ua': '"Chromium";v="118", "Google Chrome";v="118", "Not=A?Brand";v="99"',
#   'sec-ch-ua-mobile': '?0',
#   'sec-ch-ua-platform': '"Windows"',
#   'sec-fetch-dest': 'empty',
#   'sec-fetch-mode': 'cors',
#   'sec-fetch-site': 'cross-site',
#   'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36',
#   'Cookie': '__cf_bm=tS4l0uFq.mBcXFmNX4iEIH.xZ8NBnSnXgVQ.ucwD0WY-1698051802-0-AegAfnFejxqB8DIYrLjEFY47VDq+UBNrBOE+wxNlrOHhXpMaf75OUDcHZBMIy5hTqYRW0WejySpXA13uXzLZNXs='
# }

# response = requests.request("POST", url, headers=headers, data=payload)

# #Save
# with open('article.json', 'w', encoding='utf-8') as f:
#     json.dump(response.text, f)
with open('article.json', encoding='utf-8') as f:
    out = json.load(f)



link_fail = []
link_done = []

n_item = 0
for i in out['data']['articleList']['items']:
    if i['link'] not in done_link:
        print(f"SCRAPE ITEM {n_item}")
        n_item += 1
        try:
            url = 'https://www.rumah.com/panduan-properti' + i['link']
            driver.get(url=url)

            soup = BeautifulSoup(driver.page_source, 'html.parser')

            page = soup.find('div',{'class':'ArticlePageTemplate_articlePageContainer__7LtlA container'})
            sec = page.find_all('div',{'class':'row'})

            document = Document()

            article_head = sec[0].text
            article_aut = sec[1].text
            document.add_heading(article_head, 0)
            document.add_heading(article_aut, level=1)

            article_img = sec[2].find('img').get('src')
            with open('temp_img/file.png', 'wb') as f:
                f.write(requests.get(article_img).content)

            document.add_picture('temp_img/file.png',width=Inches(5))

            article_content = sec[-1].find('div',{'class':'content-wrapper'})
            content_items = article_content.findChildren()

            for item in content_items:
                to_doc(item)

            filename = re.sub(r'[^\w]', ' ', article_head)
            foldername = i['category'][0]['name']
            # folder = os.path.join('output_docs/', i['category']['name'])
            os.makedirs(f"output_docs/{str(foldername)}", exist_ok=True)
            document.save(f"output_docs/{foldername}/{filename}.docx")

            done = {'link': i['link']}
            link_done.append(done)
            with open('output_docs/done_link_3.json', 'w') as f:
                json.dump(link_done, f)

        except:
            fail = {'link': i['link']}
            link_fail.append(fail)
            with open('output_docs/fail_link_3.json', 'w') as f:
                json.dump(link_fail, f)