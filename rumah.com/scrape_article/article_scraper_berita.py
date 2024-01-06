from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
import time
from bs4 import BeautifulSoup
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
import json
from docx import Document
from htmldocx import HtmlToDocx
import requests
from docx.shared import Inches
import undetected_chromedriver as uc 
import re
import os


def to_doc(x):

    if x.name == 'p':
        document.add_paragraph(x.text)
    elif x.name == 'ol':
        li = x.find_all('li')
        for it in li:
            document.add_paragraph(it.text,style='List Bullet')
    




chrome_options = uc.ChromeOptions() 
chrome_options.headless = False  # Set headless to False to run in non-headless mode

driver = uc.Chrome(options=chrome_options,use_subprocess=True)
link_done = []
link_fail = []
year = ['2022','2023']
art_item = 0
for page in year: # LOOP YEAR
    
    url = 'https://www.rumah.com/berita-properti/archive/' + page 
    driver.get(url=url)
    #GET LAST PAGE
    soup = BeautifulSoup(driver.page_source, 'html.parser')
    try:
        last_page = soup.find('a',string='Terakhir').get('href').split('=')[-1]
        last_page = int(last_page)
        print(last_page)
    except:
        continue

    
    #LOOP THROUGH PAGES
    for i in range(1,last_page+1):
        url_page = f'https://www.rumah.com/berita-properti/archive/{page}?page={str(i)}'
        driver.get(url_page)
        time.sleep(1)
        soup_page = BeautifulSoup(driver.page_source, 'html.parser')

        
        items_wrap = soup_page.find('ul',{'class':'lists'})
        items = items_wrap.find_all('a') # THIS IS SHOULD BE LIST OF URLS
        items = [link.get('href') for link in items]

        
        for j in items: #SCRAPE ARTICLES
            try :
                art_item += 1
                print(f"SCRAPE ITEM {art_item}")
                url_article = 'https://www.rumah.com' + j
                driver.get(url_article)
                soup_article = BeautifulSoup(driver.page_source, 'html.parser')

                article_head = soup_article.find('h1',{'class':'title-big'}).text
                article_info = soup_article.find('p',{'class':'news-quick-info'}).text
                article_img = soup_article.find('img',{'class':'imgborder mleft10 mbottom10'}).get('src')
         

                with open('temp_img/file2.jpg', 'wb') as f:
                    f.write(requests.get(article_img).content)

                document = Document()
                document.add_heading(article_head, 0)
                document.add_heading(article_info, level=1)
                try:
                    document.add_picture('temp_img/file2.jpg',width=Inches(5))
                except:
                    pass

                article_body = soup_article.find('div',{'class':'news-article-body'}).findChildren()
                for par in article_body:
                    to_doc(par)


                filename = re.sub(r'[^\w]', ' ', article_head)[:30]
                foldername = j.split('/')[2]

                os.makedirs(f"berita_properti/{foldername}", exist_ok=True)
                document.save(f"berita_properti/{foldername}/{filename}.docx")


                done = {'link': url_article}
                link_done.append(done)
                with open('berita_properti/done_link2.json', 'w') as f:
                    json.dump(link_done, f)
            except :
                url_article = 'https://www.rumah.com' + j
                fail = {'link': url_article}
                link_fail.append(fail)
                with open('berita_properti/fail_link2.json', 'w') as f:
                    json.dump(link_fail, f)


