from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
import time
from bs4 import BeautifulSoup
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
import json
from docx import Document
from htmldocx import HtmlToDocx
import requests
from docx.shared import Inches
import undetected_chromedriver as uc 
import re
import os
import shutil
from PIL import Image
from io import StringIO

def to_doc(x,document):

    if x.name == 'p':
        document.add_paragraph(x.text)
    elif x.name == 'ol':
        li = x.find_all('li')
        for j in li:
            document.add_paragraph(j.text,style='List Bullet')

# def to_doc(x):

#     if x.name == 'p':
#         document.add_paragraph(x.text)

        
#     elif x.name == 'ol':
#         li = x.find_all('li')
#         for j in li:
#             document.add_paragraph(j.text,style='List Bullet')
#     elif x.name == 'h2':
#         document.add_heading(x.text, level=2)
        
#     elif x.name == 'figure':
#         try:
#             link_img = x.find('img').get('src')
#             with open('temp_img/file.png', 'wb') as f:
#                 f.write(requests.get(link_img).content)
#             document.add_picture('temp_img/file.png',width=Inches(5))

#         except:
#             pass

    



def main():
    chrome_options = uc.ChromeOptions() 
    chrome_options.headless = False  # Set headless to False to run in non-headless mode

    driver = uc.Chrome(options=chrome_options)
    link_done = []
    link_fail = []
    url_home = 'https://www.rumah.com/areainsider' 
    driver.get(url=url_home)
    soup = BeautifulSoup(driver.page_source, 'html.parser')

    listing_card = soup.find_all('div',{'class':'col-md-4 col-xs-6'})
    card_links = [i.find('a').get('href') for i in listing_card]

    scraped_item = 0
    for card in card_links:
        url = 'https://www.rumah.com' + card
        driver.get(url=url)
        soup_card = BeautifulSoup(driver.page_source, 'html.parser')
        article_cards = soup_card.find_all('div',{'class':'col-sm-4'})
        article_cards_link = [i.find('a').get('href') for i in article_cards]
        for article_url in article_cards_link:
            try:
                scraped_item += 1
                print(f"SCRAPED ITEM {scraped_item}")
                document = Document()
                url_article = 'https://www.rumah.com' + article_url
                driver.get(url=url_article)

                time.sleep(1)
                soup_article = BeautifulSoup(driver.page_source, 'html.parser')

                article_head = soup_article.find('h1',{'itemprop':'headline'}).text
                

                article_info = soup_article.find('span',{'itemprop':'datePublished'}).text
                article_firstp = soup_article.find('article',{'class':'first-paragraph'}).text
                article_body = soup_article.find('div',{'class':'article-body'})
                article_body = article_body.find_all('article')[-1].findChildren()

                # print('panajnag body',len(article_body[-1].findChildren()))


                try:
                    article_img = driver.find_element(By.CLASS_NAME,"article-featured-image")
                    article_img.screenshot('temp_img/file_temp.png')
                    # img_url = soup_article.find('img',{'class':'article-featured-image'}).get('src')
                    # print(img_url)

                    # r = requests.get(img_url)

                    # i = Image.open(StringIO(r.content))
                    # i.save('temp_img/file7.jpg')
                    # with open('temp_img/file6.jpg', 'wb') as f:
                    #     f.write(article_img.screenshot_as_png)
                except:
                    pass
                document.add_heading(article_head, 0)
                document.add_paragraph(article_info)
                document.add_paragraph(article_firstp)

                try:
                    document.add_picture('temp_img/file_temp.jpg',width=Inches(5))
                except:
                    pass
                
                for par in article_body:
                    to_doc(par,document)

                filename = re.sub(r'[^\w]', ' ', article_head)
                foldername = article_url.split('/')[2]
                os.makedirs(f"areainsider/{foldername}", exist_ok=True)
                document.save(f"areainsider/{foldername}/{filename}.docx")

                done = {'link': article_url}
                link_done.append(done)
                with open('areainsider/done_link.json', 'w') as f:
                    json.dump(link_done, f)
            except:
                fail = {'link': article_url}
                link_fail.append(fail)
                with open('areainsider/fail_link.json', 'w') as f:
                    json.dump(link_fail, f)
