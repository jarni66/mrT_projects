from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
import time
from bs4 import BeautifulSoup
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
import json
from docx import Document
from htmldocx import HtmlToDocx
import requests
from docx.shared import Inches
import undetected_chromedriver as uc 
import re
import os


# def to_doc(x):

#     if x.name == 'p':
#         document.add_paragraph(x.text)
#     elif x.name == 'ol':
#         li = x.find_all('li')
#         for j in li:
#             document.add_paragraph(j.text,style='List Bullet')

def to_doc(x,document):

    if x.name == 'p':
        document.add_paragraph(x.text)

        
    elif x.name == 'ol':
        li = x.find_all('li')
        for j in li:
            document.add_paragraph(j.text,style='List Bullet')
    elif x.name == 'h2':
        document.add_heading(x.text, level=2)
        
    elif x.name == 'figure':
        try:
            link_img = x.find('img').get('src')
            with open('temp_img/file.png', 'wb') as f:
                f.write(requests.get(link_img).content)
            document.add_picture('temp_img/file.png',width=Inches(5))

        except:
            pass

    
def main():



    chrome_options = uc.ChromeOptions() 
    chrome_options.headless = False  # Set headless to False to run in non-headless mode

    driver = uc.Chrome(options=chrome_options)
    link_done = []
    link_fail = []
    categories = ['pembaruan-produk/']
    item_scrape = 0
    for category in categories: # LOOP YEAR
        
        url = 'https://agentofferings.rumah.com/kategori/' + category 
        driver.get(url=url)

        #GET LAST PAGE
        soup = BeautifulSoup(driver.page_source, 'html.parser')
        last_page = soup.find('ul',{'class':'page-numbers'})
        if last_page != None:
            last_page = last_page.find_all('li')[-1]
            last_page = int(last_page.text)

            #LOOP THROUGH PAGES
            for i in range(1,last_page+1):
                url_page = f'https://agentofferings.rumah.com/kategori/wawasan-and-tips/page/{i}/'
                driver.get(url_page)
                soup_page = BeautifulSoup(driver.page_source, 'html.parser')
                items = soup_page.find_all('h3',{'class':'pg-h3'})
                items = [k.find('a').get('href') for k in items] # THIS IS SHOULD BE LIST OF URLS

                for j in items: #SCRAPE ARTICLES
                    item_scrape += 1
                    print(f"ITEM SCRAPED {item_scrape}")
                    try :
                        document = Document()
                        driver.get(j)
                        soup_article = BeautifulSoup(driver.page_source, 'html.parser')

                        article_head = soup_article.find('h1',{'class':'pg-h1'}).text
                        article_info = soup_article.find('div',{'class':'entry-meta'}).text
                        article_img = soup_article.find('img',{'class':'attachment-post-thumbnail size-post-thumbnail wp-post-image'}).get('src')
                        with open('temp_img/file.png', 'wb') as f:
                            f.write(requests.get(article_img).content)

                        document.add_heading(article_head, 0)
                        document.add_heading(article_info, level=1)
                        try:
                            document.add_picture('temp_img/file.png',width=Inches(5))
                        except:
                            pass


                        article_body = soup_article.find('div',{'class':'entry-content'}).findChildren()
                        for par in article_body:
                            to_doc(par,document)
                        
                        filename = re.sub(r'[^\w]', ' ', article_head)

                        foldername = category
                        # folder = os.path.join('output_docs/', i['category']['name'])
                        os.makedirs(f"bantuan_properti/{foldername}", exist_ok=True)
                        document.save(f"bantuan_properti/{foldername}/{filename}.docx")

                        done = {'link': j}
                        link_done.append(done)
                        with open('bantuan_properti/done_link2.json', 'w') as f:
                            json.dump(link_done, f)
                    except :
                        fail = {'link': j}
                        link_fail.append(fail)
                        with open('bantuan_properti/fail_link2.json', 'w') as f:
                            json.dump(link_fail, f)

        else :
            soup_page = BeautifulSoup(driver.page_source, 'html.parser')
            items = soup_page.find_all('h3',{'class':'pg-h3'})
            items = [k.find('a').get('href') for k in items] # THIS IS SHOULD BE LIST OF URLS

            for j in items: #SCRAPE ARTICLES
                item_scrape += 1
                print(f"ITEM SCRAPED {item_scrape}")
                try :
                    document = Document()
                    driver.get(j)
                    soup_article = BeautifulSoup(driver.page_source, 'html.parser')

                    article_head = soup_article.find('h1',{'class':'pg-h1'}).text
                    article_info = soup_article.find('div',{'class':'entry-meta'}).text
                    article_img = soup_article.find('img',{'class':'attachment-post-thumbnail size-post-thumbnail wp-post-image'}).get('src')
                    with open('temp_img/file.png', 'wb') as f:
                        f.write(requests.get(article_img).content)

                    document.add_heading(article_head, 0)
                    document.add_heading(article_info, level=1)
                    try:
                        document.add_picture('temp_img/file.png',width=Inches(5))
                    except:
                        pass


                    article_body = soup_article.find('div',{'class':'entry-content'}).findChildren()
                    for par in article_body:
                        to_doc(par)
                    
                    filename = re.sub(r'[^\w]', ' ', article_head)

                    foldername = category
                    # folder = os.path.join('output_docs/', i['category']['name'])
                    os.makedirs(f"bantuan_properti/{foldername}", exist_ok=True)
                    document.save(f"bantuan_properti/{foldername}/{filename}.docx")

                    done = {'link': j}
                    link_done.append(done)
                    with open('bantuan_properti/done_link2.json', 'w') as f:
                        json.dump(link_done, f)
                except :
                    fail = {'link': j}
                    link_fail.append(fail)
                    with open('bantuan_properti/fail_link2.json', 'w') as f:
                        json.dump(link_fail, f)
