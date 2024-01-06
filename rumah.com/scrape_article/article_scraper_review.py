from selenium import webdriver
from selenium.webdriver.common.action_chains import ActionChains
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.options import Options
import time
from bs4 import BeautifulSoup
from selenium.webdriver.common.by import By
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import WebDriverWait
import json
from docx import Document
from htmldocx import HtmlToDocx
import requests
from docx.shared import Inches
import undetected_chromedriver as uc 
import re
import os
import shutil
from PIL import Image
from io import StringIO

# def to_doc(x):

#     if x.name == 'p':
#         document.add_paragraph(x.text)
#     elif x.name == 'ol':
#         li = x.find_all('li')
#         for j in li:
#             document.add_paragraph(j.text,style='List Bullet')

# def to_doc(x):

#     if x.name == 'p':
#         document.add_paragraph(x.text)

        
#     elif x.name == 'ol':
#         li = x.find_all('li')
#         for j in li:
#             document.add_paragraph(j.text,style='List Bullet')
#     elif x.name == 'h2':
#         document.add_heading(x.text, level=2)
        
#     elif x.name == 'figure':
#         try:
#             link_img = x.find('img').get('src')
#             with open('temp_img/file.png', 'wb') as f:
#                 f.write(requests.get(link_img).content)
#             document.add_picture('temp_img/file.png',width=Inches(5))

#         except:
#             pass

    
def scrape_article(soup_article):
    # driver.get(url=x)
    # soup_article = BeautifulSoup(driver.page_source, 'html.parser')
    article_head = soup_article.find('h1').text.strip()
    article_aut = soup_article.find('ul',{'class':'author'}).text
    article_intro = soup_article.find('div',{'class':'introduction'}).findChildren()
    article_body = soup_article.find('div',{'class':'review-detail-collapse'}).findChildren()

    
    document.add_heading(article_head, 0)
    document.add_paragraph(article_aut)
    for par in article_intro:
        document.add_paragraph(par.text)
    for par in article_body:
        document.add_paragraph(par.text)

    filename = re.sub(r'[^\w]', ' ', article_head)
    # foldername = article_url.split('/')[2]
    # os.makedirs(f"areainsider/{foldername}", exist_ok=True)
    document.save(f"review/docs/{filename}.docx")



chrome_options = uc.ChromeOptions() 
chrome_options.headless = False  # Set headless to False to run in non-headless mode

driver = uc.Chrome(options=chrome_options)
link_done = []
link_fail = []
url_home = 'https://www.rumah.com/perumahan-baru/review' 
driver.get(url=url_home)
soup = BeautifulSoup(driver.page_source, 'html.parser')
last_page = soup.find('ul',{'class':'pagination'}).find_all('a')
last_page = int(last_page[-2].text)



for page in range(1,last_page+1):
    url_page = f'https://www.rumah.com/perumahan-baru/review/{page}'
    driver.get(url=url_page)
    soup_card = BeautifulSoup(driver.page_source, 'html.parser')
    article_cards = soup_card.find_all('li',{'class':'listing-item'})
    article_cards_link = [i.find('a').get('href') for i in article_cards]
    item = 0
    # print(article_cards_link[0])
    for article_link in article_cards_link:
        article_url = 'https://www.rumah.com' + article_link
        item += 1
        print(f"SCRAPE page {page} : item {item}")
        try:
            driver.get(url=article_url)
            soup_article = BeautifulSoup(driver.page_source, 'html.parser')
            document = Document()
            scrape_article(soup_article)
            done = {'link': article_url}
            link_done.append(done)
            with open('review/done_link.json', 'w') as f:
                json.dump(link_done, f)
        except:
            print(f"FAIL page {page} : item {item}")
            fail = {'link': article_url}
            link_fail.append(fail)
            with open('review/fail_link.json', 'w') as f:
                json.dump(link_fail, f)
            
        

